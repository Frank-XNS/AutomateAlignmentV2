"""
Creating .lab (or .txt) files by segmenting a .docx file and according to given .wav files
"""

import os
import docx
import codecs

#  folder_exists(), file_exists() and no_forbidden_characters are re-usable but minor changes
#  must be brought to the error messages
FRENCH_NUMBERS_DIGITS = {0: "", 1: "un", 2: "deux", 3: "trois", 4: "quatre", 5: "cinq", 6: "six", 7: "sept", 8: "huit", 9: "neuf", 10: "dix", 11: "onze", 12: "douze", 13: "treize", 14: "quatorze", 15: "quinze", 16: "seize", 17: "dix-sept", 18: "dix-huit", 19: "dix-neuf"}
FRENCH_NUMBERS_TENS = {2: "vingt", 3: "trente", 4: "quarante", 5: "cinquante", 6: "soixante", 7: "soixante", 8: "quatre-vingt", 9: "quatre-vingt"}
FRENCH_NUMBERS_HUNDREDS = {1: "cent", 10: "mille"}
def folder_exists(this_folder: str) -> bool:
    """Check whether this_folder exists"""
    if not os.path.isdir(this_folder):
        print("\n")
        print("{} does not exist.".format(this_folder))
        print("Please make sure you typed the name of the folder containing the wav files correctly and that you are"
              " in the correct directory.\n")
        return False
    else:
        return True


def file_exists(this_file: str) -> bool:
    """Check whether this_file exists"""
    if not os.path.isfile(this_file):
        print("\n")
        print("{} does not exist.".format(this_file))
        print("Please make sure you typed the name of the docx file correctly and that you are in the correct"
              " directory.\n")
        return False
    else:
        return True


def no_forbidden_characters(this_new_item: str) -> bool:
    """Check whether this_new_item (file or folder) contains no forbidden characters"""
    forbidden_characters = ['\\', '/', ':', '*', '?', '"', '<', '>', '|']
    folder_name_wrong = any([char in forbidden_characters for char in this_new_item])
    if folder_name_wrong:
        print("You cannot use the following characters in your folder name:\n")
        print(', '.join(forbidden_characters) + '\n')
        return False
    else:
        return True


#  folder_already exists() are re-usable with no changes necessary
def folder_already_exists(this_new_folder: str)-> bool:
    """Check whether a folder called this_folder already exists"""
    this_new_folder_already_exists = os.path.isdir(this_new_folder)
    if this_new_folder_already_exists:
        print("The folder '{}' already exists. Please choose another name or delete the existing folder.\n".format(
            this_new_folder))
        return True
    else:
        return False


def convert_two_digit_number_to_french(digits):
    temp_countainer = ["",  "-"]
    # If the integer is single digit
    if int(digits[0]) < 2:
        return FRENCH_NUMBERS_DIGITS[int(digits)]
    # If the integer is double digit less than 70
    elif int(digits[0]) < 7:
        if int(digits[1]) != 1:
            return  FRENCH_NUMBERS_TENS[int(digits[0])] + ["",  "-"][int(digits[1]) != 0] + FRENCH_NUMBERS_DIGITS[int(digits[1])]
        else:
            return FRENCH_NUMBERS_TENS[int(digits[0])] + "-" + "et-un"
    # If the integer between 70 and 80
    elif int(digits[0]) < 8:
        if int(digits[1]) != 1:
            return FRENCH_NUMBERS_TENS[int(digits[0])] + ["",  "-"][int(digits[1]) != 0] + FRENCH_NUMBERS_DIGITS[
                int(digits[1]) + 10]
        else:
            return FRENCH_NUMBERS_TENS[int(digits[0])] + "-" + "et-onze"
    # If the integer is between 80 and 90
    elif int(digits[0]) < 9:
        if (int(digits[1]) != 0):
            return FRENCH_NUMBERS_TENS[int(digits[0])] + ["",  "-"][int(digits[1]) != 0] + FRENCH_NUMBERS_DIGITS[
                int(digits[1])]
        else:
            return FRENCH_NUMBERS_TENS[int(digits[0])] + "s"
    # If the integer is between 90 and 100
    elif int(digits[0]) <= 9:
        return FRENCH_NUMBERS_TENS[int(digits[0])] + ["",  "-"][int(digits[1]) != 0] + FRENCH_NUMBERS_DIGITS[
            int(digits[1]) + 10]

def generate_lab_files(word_file_name: str, folder_wav_files: str, new_folder_lab_files: str):
    """
    Look at word_file_name and generate the .lab files in new_folder_lab_files
    """

    # Both word_file_name and folder_wav_files exist, there's no forbidden character in new_folder_lab_files and
    # no other folder is alreay called new_folder_lab_files
    if (file_exists(word_file_name) and
            folder_exists(folder_wav_files) and
            no_forbidden_characters(new_folder_lab_files) and
            not folder_already_exists(new_folder_lab_files)):

        # Open the given word document
        doc = docx.Document(word_file_name)

        # Let's consider an empty paragraph contains an empty string or a succession of spaces (here up to 9 spaces)
        # We can make that list bigger if I haven't been able to predict all possible types of meaningless paragraphs
        list_empty_para = [' ' * x for x in range(10)]

        # Extract the text into a list and ignore empty paragraphs (these tend to be at the end of doc)
        text_doc_raw = [doc.paragraphs[i].text for i in range(len(doc.paragraphs))
                    if doc.paragraphs[i].text not in list_empty_para]
        num_para = len(text_doc_raw)
        print("There are %(num)d paragraphs in total\n" %{'num': num_para})
        #Removes all comma and :
        for paragraph in range(len(text_doc_raw)):
            thisParagraph = list(text_doc_raw[paragraph])
            for index in range(len(thisParagraph)):
                if thisParagraph[index] == ',' or thisParagraph[index] == ':' or thisParagraph[index] == '.':
                    thisParagraph[index] = ''
                if thisParagraph[index] == '-' or thisParagraph[index] == "'" or thisParagraph[index] == "’":
                    thisParagraph[index] = ' '
            text_doc_raw[paragraph] = ''.join(thisParagraph)
        #Converts all integers to french words

        text_doc = []
        resolved_number = {}

        #This converts the numbers into digits
        for paragraph in text_doc_raw:
            paragraph_temp = paragraph.split(" ")
            for index in range(len(paragraph_temp)):
                has_h_at_end = 0
                if paragraph_temp[index].endswith(','):
                    paragraph_temp[index] = paragraph_temp[index][:-1]
                if paragraph_temp[index].endswith('h'):
                    has_h_at_end = 1
                    paragraph_temp[index] = paragraph_temp[index][:-1]
                if paragraph_temp[index].isdigit():
                    #For increasing integer cases
                    if (int(paragraph_temp[index]) < 10):
                        paragraph_temp[index] = convert_two_digit_number_to_french('0' + paragraph_temp[index])
                    elif (int(paragraph_temp[index]) < 100):
                        paragraph_temp[index] = convert_two_digit_number_to_french(paragraph_temp[index])
                    elif (int(paragraph_temp[index]) < 1000):
                        paragraph_temp[index] = FRENCH_NUMBERS_DIGITS[int(paragraph_temp[index][0])] + " cent " + convert_two_digit_number_to_french(paragraph_temp[index][1:3])
                    elif (int(paragraph_temp[index]) < 10000):
                        if int(paragraph_temp[index]) in resolved_number:
                            paragraph_temp[index] = resolved_number[int(paragraph_temp[index])]
                        else:
                            number = int(paragraph_temp[index])
                            #case 1: years as vigésimales system
                            caseOne = convert_two_digit_number_to_french(paragraph_temp[index][0:2]) + " cent " + convert_two_digit_number_to_french(paragraph_temp[index][2:4])
                            #case 2: years as regular number
                            caseTwo = FRENCH_NUMBERS_DIGITS[int(paragraph_temp[index][0])] + " mille " + FRENCH_NUMBERS_DIGITS[int(paragraph_temp[index][1])] + [""," cent "][int(paragraph_temp[index][1]) != 0] + convert_two_digit_number_to_french(paragraph_temp[index][2:4])
                            sel = input("How is the year %(year)s pronounced?\nFor %(caseOne)s (vigésimales system) press 1 and enter\nFor %(caseTwo)s (regular number) press 2 and enter\n" % {'year': paragraph_temp[index], 'caseOne':caseOne, 'caseTwo': caseTwo})
                            if sel == 1:
                                paragraph_temp[index] = caseOne
                            else:
                                paragraph_temp[index] = caseTwo
                            resolved_number[number] = paragraph_temp[index]
                    else:
                        #For a string of digits
                        if int(paragraph_temp[index]) in resolved_number:
                            paragraph_temp[index] = resolved_number[int(paragraph_temp[index])]
                        else:
                            number = int(paragraph_temp[index])
                            digit_word = []
                            for digit in paragraph_temp[index]:
                                if int(digit) == 0:
                                    digit_word.append("zéro")
                                digit_word.append(FRENCH_NUMBERS_DIGITS[int(digit)])
                            sel_large_number = input("The script has generated %(guessed)s for the number %(number)s, if you want to use this,\npress 0 and enter, else input the translation and enter"
                                                     %{"guessed": " ".join(digit_word), "number": number})
                            if sel_large_number == "0":
                                paragraph_temp[index] = " ".join(digit_word)
                            else:
                                paragraph_temp[index] = sel_large_number
                            resolved_number[number] = paragraph_temp[index]
                    if has_h_at_end == 1:
                        paragraph_temp[index] = paragraph_temp[index] + ' heure'

            text_doc.append(" ".join(paragraph_temp))

        # Get the names of the wav files
        wav_files = os.listdir(folder_wav_files)
        num_wav = len(wav_files)

        if num_para != num_wav:
            print("\nThere are {} wav files in '{}' but {} paragraphs in '{}'.".format(num_wav, folder_wav_files,
                                                                                       num_para, word_file_name))
            print("The number of wav files must be the same as the number of paragraphs.")
            print("No lab file has been created.")
            print("Please adjust your files/folders accordingly before re-running the code.\n")

        else:
            # Create the folder for the lab files and move into it
            os.mkdir(new_folder_lab_files)
            os.chdir(new_folder_lab_files)
            sel_wave_file_sorted_order = input("Is the wave files arranged in ascending name?(y/n)")
            if sel_wave_file_sorted_order == 'y':
                wav_files.sort()
            else:
                wav_files.sort(reverse = True)
            for i in range(num_wav):
                #Remove the .wav extention and replace with with .lab
                lab_file_name = wav_files[i][:-4] + ".lab"

                lab_file_content = text_doc[i].replace("’", "'").replace("œ", "oe")

                # lab_file = open(lab_file_name, 'w')
                # lab_file.write(lab_file_content)
                with codecs.open(lab_file_name, 'w', encoding="utf-8") as lab_file:
                    lab_file.write(lab_file_content)

            print("Success!\n")


if __name__ == '__main__':
    print("The paragraphs in the docx file must match the recordings in the folder containing the wav files.")
    word_file_name = input("Enter the name of the docx file: ")
    folder_wav_files = input("Enter the name of the folder containing the wav files: ")
    new_folder_lab_files = input("Enter the name for a new folder to contain the lab files: ")
    generate_lab_files(word_file_name + ".docx", folder_wav_files, new_folder_lab_files)
